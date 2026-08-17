from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
out = ROOT / "reference.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei UI"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(23, 43, 77)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.35

for name, size, color in [
    ("Title", 30, "102A56"),
    ("Subtitle", 16, "246BFD"),
    ("Heading 1", 22, "102A56"),
    ("Heading 2", 16, "173D72"),
    ("Heading 3", 12, "24466F"),
]:
    style = styles[name]
    style.font.name = "Microsoft YaHei UI"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)

styles["Title"].paragraph_format.space_after = Pt(10)
styles["Heading 1"].paragraph_format.space_before = Pt(18)
styles["Heading 1"].paragraph_format.space_after = Pt(8)
styles["Heading 2"].paragraph_format.space_before = Pt(14)
styles["Heading 2"].paragraph_format.space_after = Pt(6)

for style_name in ("Table", "Compact"):
    if style_name in styles:
        styles[style_name].font.name = "Microsoft YaHei UI"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
        styles[style_name].font.size = Pt(9)

doc.core_properties.title = "RoadMan｜多智能体自驾与旅行规划工作台"
doc.core_properties.subject = "GOAI 无界应用｜Boundless Agents · AI+汽车参赛方案"
doc.core_properties.author = "RoadMan Team"

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("RoadMan · GOAI 无界应用参赛方案")
run.font.name = "Microsoft YaHei UI"
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(127, 145, 170)

doc.save(out)
print(out)
